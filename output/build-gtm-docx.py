#!/usr/bin/env python3
"""
Build output/Sealshot-Go-To-Market-Plan.docx.

The two-stage launch plan, with the corrections from the 3 August review folded
in — most importantly Gate 0, which makes the original timeline impossible as
written. US English throughout.

    python3 output/build-gtm-docx.py
"""
import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0xC2, 0x41, 0x0C)
INK = RGBColor(0x0B, 0x12, 0x20)
MUTED = RGBColor(0x4F, 0x5A, 0x72)
RED = RGBColor(0x9A, 0x34, 0x12)

OUT = pathlib.Path(__file__).resolve().parent / 'Sealshot-Go-To-Market-Plan.docx'

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Helvetica Neue'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before in [
    ('Heading 1', 20, INK, 22), ('Heading 2', 14, ACCENT, 18), ('Heading 3', 11.5, INK, 13),
]:
    st = doc.styles[name]
    st.font.name = 'Helvetica Neue'
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True

for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    sec.left_margin = sec.right_margin = Inches(0.9)


def para(text='', style=None, bold=False, italic=False, color=None, size=None, after=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = Pt(size)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def rich(*parts, style=None, after=None):
    p = doc.add_paragraph(style=style)
    for part in parts:
        text, opts = part if isinstance(part, tuple) else (part, {})
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        r.italic = opts.get('italic', False)
        if 'color' in opts:
            r.font.color.rgb = opts['color']
        if opts.get('mono'):
            r.font.name = 'Menlo'
            r.font.size = Pt(9.5)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def bullet(*parts, num=False):
    p = doc.add_paragraph(style='List Number' if num else 'List Bullet')
    p.paragraph_format.space_after = Pt(3)
    for part in parts:
        text, opts = part if isinstance(part, tuple) else (part, {})
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        r.italic = opts.get('italic', False)
        if opts.get('mono'):
            r.font.name = 'Menlo'
            r.font.size = Pt(9.5)
    return p


def shade(cell, fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(el)


def table(headers, rows, widths=None, bold_first=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        shade(c, 'F1F4F9')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if bold_first and i == 0:
                r.bold = True
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def callout(title, body_paras, fill='FDF0E8'):
    """A boxed warning — used for the gates that stop a launch."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    shade(cell, fill)
    cell.width = Inches(6.5)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RED
    for b in body_paras:
        q = cell.add_paragraph()
        q.paragraph_format.space_after = Pt(5)
        q.paragraph_format.line_spacing = 1.15
        rr = q.add_run(b)
        rr.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    return p


# ══════════════════════════════════════════════════════════════════ title ════
para('Sealshot', style='Heading 1', after=0)
para('Go-to-market plan', bold=True, color=ACCENT, size=13, after=2)
para('Two-stage launch · reviewed 3 August 2026', italic=True, color=MUTED, after=14)

para('A small paid Founding release first, then the broad 1.0 launch once there '
     'are customers, testimonials, and a message that has been tested on '
     'strangers.')

rich(('The central strategic choice: ', {'bold': True}),
     ('market Sealshot as the safest way to share what is on your screen — not '
      'as the screenshot app with the longest feature list. Capture, recording, '
      'OCR, encryption and local AI then become supporting evidence for that '
      'promise rather than competing headlines.', {}))

# ══════════════════════════════════════════════════════════════════ gate 0 ═══
para('Gate 0 — before a single sale or public claim', style='Heading 2')

para('Three things sit ahead of everything else. Each is currently unverified, '
     'and each would be discovered by a customer or by Hacker News rather than '
     'by us. They are tracked in docs/launch-checklist.md.')

callout('G0.1 — The app cannot yet accept what we sell', [
    'The fulfillment Worker emits license preamble v2. The newest tagged release '
    'is v0.7.2, which verifies v1. A purchase made today produces a file that '
    'the app the customer just downloaded rejects as tampered — they pay and '
    'cannot activate.',
    'A second reason the release must come first: a founding buyer\'s 18-month '
    'window is measured against the build\'s entitlement date, stamped at '
    'release time. Sell before the release exists and the arithmetic has nothing '
    'to measure against.',
    'This makes the original "run the founding cohort now" timeline impossible. '
    'Cut a Direct release containing licensing v2, verify a Worker-issued '
    'license activates on the build downloaded from its public URL, then sell.',
])

callout('G0.2 — Redaction permanence is unproven', [
    'The entire message is "share screenshots without sharing secrets". '
    'Redaction permanence is Tier 1 and unverified, with a specific suspected '
    'failure: the exported flat PNG is clean while the layered .seal package '
    'still carries the original underneath — so sharing the package leaks what '
    'sharing the PNG does not.',
    'If someone recovers redacted content after launch, that is not a bug '
    'report. It ends the positioning permanently. Attack your own export before '
    'building a launch on the claim.',
])

callout('G0.3 — Network silence is unproven', [
    '"No telemetry" is the core claim and the privacy policy commits to it in '
    'writing. It has never been verified at the packet level.',
    'Watch a full session through Little Snitch or a proxy. Then consider '
    'publishing the result — almost no competitor can, and Show HN readers will '
    'run this test whether or not you do.',
])

# ═════════════════════════════════════════════════════════════════ 1 beachhead
para('1. A narrow beachhead', style='Heading 2')

quote('Mac-based developers, QA engineers, and support professionals who '
      'regularly share screenshots containing customer data, credentials, admin '
      'panels, or internal systems.')

para('They are reachable through founder-led outreach, Mac communities, Hacker '
     'News and Product Hunt — and, more importantly, they already have the '
     'problem: screenshots leak secrets.')

para('Avoid positioning for every Mac user, and avoid healthcare or legal '
     'enterprises initially. Those need credibility you do not have yet and '
     'sales cycles you cannot fund.')

para('Positioning statement', style='Heading 3')
quote('Sealshot helps Mac professionals share screenshots without sharing '
      'secrets. It detects and permanently redacts sensitive information '
      'on-device — with no account, no telemetry, and no cloud upload.')

para('Homepage headline', style='Heading 3')
para('Replace the current feature-oriented headline in src/pages/index.astro '
     'with an outcome-oriented one:')
quote('Share screenshots without sharing secrets. Capture, annotate, and find '
      'sensitive information before you send it — all on your Mac. No account. '
      'No telemetry. No cloud.')
rich('Primary CTA ', ('Start your free 14-day trial', {'bold': True}),
     ', secondary ', ('See Smart Redaction', {'bold': True}), '.')

# ═══════════════════════════════════════════════════════════════════ 2 pricing
para('2. Keep the pricing', style='Heading 2')

para('Verified against the live competition on 3 August 2026:')

table(['Product', 'Model', 'Account required'],
      [['CleanShot X', '$29 one Mac, $49 two Macs, $19/yr optional updates. 30-day money-back guarantee', 'No'],
       ['Snagit', '$39 per user per year, subscription, non-transferable', 'Yes — "just sign in"'],
       ['Shottr', '$12 one-time', 'No'],
       ['Sealshot', '$29.99 perpetual, three Macs, updates optional', 'None']],
      widths=[1.1, 3.9, 1.5], bold_first=True)

rich(('Snagit is the comparison being under-used. ', {'bold': True}),
     ('It is the mainstream incumbent, and it charges $39 per user every year, '
      'with an account required to unlock the software. Sealshot is $29.99 once '
      '— so Snagit costs more in its first year than Sealshot costs ever, and '
      'more again every year after. "Perpetual, no account, nothing leaves your '
      'Mac" against "$39/year, sign in to unlock" is the sharpest single line '
      'available, and the positioning statement does not reach for it.', {}))

para('Shottr owns the cheap end at $12, and the price cut to $29.99 narrows '
     'that gap without closing it — so competing on price is still both '
     'difficult and strategically weak. Sealshot earns its premium through '
     'privacy, encryption, automatic redaction, offline licensing and a '
     'perpetual license.')

para('Keep', style='Heading 3')
bullet(('Founding: ', {'bold': True}), '$14.99, three Macs, 18 months of updates.')
bullet(('1.0 onward: ', {'bold': True}), '$29.99, three Macs, 12 months of updates.')
bullet(('Renewal: ', {'bold': True}), '$17.99, optional, another 12 months — 40% off a new license.')

para('Add no further discounts. End the founding offer on the actual 1.0 date — '
     'the endsAt in src/config/promos.ts is still a placeholder, and selling '
     '"founding" after 1.0 ships means charging $14.99 for a tier that no longer '
     'exists.')

para('Two gaps on the buy page, both now closed', style='Heading 3')
bullet('"One user, three Macs, commercial use included" was buried in the FAQ and '
       'is now on the buy page, where a buyer looks for terms.')
bullet('Refund terms did not exist. CleanShot publishes a 30-day guarantee at '
       'the same price; /refunds now states the same, linked from the buy page '
       'and the footer.')

# ════════════════════════════════════════════════════════════════ 3 founding ══
para('3. The paid founding cohort — 3 to 4 weeks', style='Heading 2')

para('Recruit 20–30 people by hand:')
bullet('5 complimentary design partners, chosen for feedback quality.')
bullet('15–25 paying the real $14.99 founding price. At the new price that is '
       '$225–375 in cohort revenue rather than $585–975 — the cohort is for '
       'evidence that people will pay at all, not for money.')
bullet('At least five buyers who are not friends or close contacts.')

para('Recruit from Mac developers and indie app makers, technical support and '
     'customer-success people, security consultants, privacy-conscious small '
     'SaaS teams, and existing contacts who routinely send bug reports or '
     'customer screenshots.')

rich(('Before any external sale: buy your own product five times. ', {'bold': True}),
     ('Clean machines, five different email addresses, the real price. The '
      'design partners are free, so without this the first exercise of the paid '
      'path is a stranger\'s money.', {}))

para('Give each person one concrete assignment — a real bug report, support '
     'reply or sensitive document — not a generic tour of the app.')

para('Interview at day seven', style='Heading 3')
for q in ['What were you using before?',
          'What real screenshot did you trust Sealshot with?',
          'Which feature made you install it?',
          'What nearly made you stop using it?',
          'Would you recommend it, and to whom?']:
    bullet(q, num=True)

para('Graduate to 1.0 when', style='Heading 3')
bullet('10+ paid buyers.')
bullet('Five reporting repeat use.')
bullet('Three specific testimonials, and one detailed customer workflow.')
bullet('No unresolved purchase, data-loss, redaction or export failures.')
bullet('20 consecutive successful license deliveries — countable as order '
       'records in KV with state "sent".')

# ══════════════════════════════════════════════════════════ 4 demonstration ═══
para('4. Build the launch around one demonstration', style='Heading 2')

para('A 30–45 second video is the primary asset:')
for step in ['Capture an admin panel containing an API key and an email address.',
             'Run Smart Redaction.',
             'Review the detections and cover them permanently.',
             'Annotate the actual problem.',
             'Export the safe image.',
             'Close on: "Everything happened on this Mac."']:
    bullet(step, num=True)

rich(('That closing line is currently an assertion. ', {'bold': True}),
     ('G0.3 is what makes it a statement of fact.', {}))

para('Three short clips cut from it: "Catch an API key before it reaches the '
     'ticket", "Extract a table from a screenshot without uploading it", '
     '"Encrypt your screenshot library behind Touch ID".')

para('Also before 1.0', style='Heading 3')
bullet('An honest comparison page against macOS, Shottr, CleanShot and Snagit.')
bullet('A concise security and network-behavior page — ideally carrying G0.3\'s result.')
bullet('A press kit: icon, screenshots, demo, pricing, requirements, founder bio.')
bullet('Testimonials beside the download CTA.')
bullet('A volume-license CTA. The FAQ promises business licensing; the buy page '
       'sells only the individual product, and licensing@ needs to reach a '
       'person before a launch drives the largest order available.')

# ═════════════════════════════════════════════════════════════ 5 distribution ═
para('5. Sequence distribution', style='Heading 2')

para('One channel at a time, so each can be engaged with properly.')

table(['#', 'Channel', 'Note'],
      [['1', 'Founding cohort', 'Direct outreach and personal communities'],
       ['2', 'Mac community', 'Participate in r/macapps before promoting; direct-download apps may be restricted to the App Pile megathread — state problem, comparison and price explicitly'],
       ['3', 'Show HN', 'Lead with the technical and privacy story; be present for architecture questions'],
       ['4', 'Product Hunt', 'Reserve for 1.0, with testimonials and a polished demo. Launch it yourself; never ask directly for upvotes'],
       ['5', 'Review outreach', 'Personal pitches and licenses to a handful of Mac writers and YouTubers'],
       ['6', 'Search content', 'High-intent articles — see below']],
      widths=[0.3, 1.5, 4.7])

para('Article targets: "How to redact a screenshot safely on Mac", "Why '
     'pixelation is not secure redaction", "Private screenshot tools for Mac", '
     '"CleanShot alternative with no cloud or account", "How to remove API keys '
     'from bug-report screenshots".')

rich(('No paid ads yet. ', {'bold': True}),
     ('A $29.99 perpetual product cannot absorb much acquisition cost before you '
      'understand conversion and renewal behavior — and at $14.99 founding, a '
      'paid click costing more than a few dollars needs a conversion rate no new '
      'product has evidence for.', {}))

# ═══════════════════════════════════════════════════════════════ 6 measurement
para('6. Measure the business without compromising the product', style='Heading 2')

para('Cloudflare Web Analytics for traffic, Polar for purchase attribution. '
     'Polar checkout links accept UTM parameters and attach them to the '
     'checkout session metadata, which carries onto the resulting order — so '
     'source attribution survives to the sale.')

para('Track weekly: visitors by channel; download-page visits and GitHub asset '
     'downloads; buy-page visits; checkout starts and completions; revenue, '
     'refunds and support volume; purchases by UTM source; interview-reported '
     'repeat use.')

table(['Planning threshold', 'Target'],
      [['Qualified visitors reaching the download page', '15%'],
       ['Buy-page visitors starting checkout', '20%'],
       ['Started checkouts completing', '50%'],
       ['Overall visitor to paid', '1%'],
       ['Refunds', 'under 5%']],
      widths=[4.6, 1.9])

para('These diagnose the funnel. They are not universal benchmarks.')

callout('The funnel above cannot see where the decision happens', [
    'There is no account, and no email is captured anywhere in the app. That '
    'means a trial user cannot be contacted — no day-3 tip, no day-12 reminder, '
    'no win-back. The trial banner appears only at seven days remaining or '
    'fewer.',
    'So the entire conversion mechanism is an in-app banner in the final week, '
    'and a buyer who trials for two weeks then purchases appears in analytics as '
    'an unattributed buy-page visit.',
    'Two consequences. The highest-leverage conversion work is the day 7 to 14 '
    'in-app experience, not the landing page. And the newsletter is the only '
    'audience that can be contacted twice — which makes the ask at trial start '
    'matter far more than it looks.',
], fill='F1F4F9')

# ══════════════════════════════════════════════════════════════════ sequence ══
para('Revised sequence', style='Heading 2')

for step in [
    'Ship a Direct release containing licensing v2 (G0.1).',
    'Verify redaction permanence and network silence (G0.2, G0.3) — the two '
    'claims the launch rests on.',
    'Production Polar, mail.seal-shot.com verified in Resend, refund terms '
    'published, licensing@ monitored.',
    'Five self-purchases on clean machines.',
    'Founding cohort, 3–4 weeks.',
    'Sequenced distribution.',
]:
    bullet(step, num=True)

para('The renewal cliff', style='Heading 3')
para('An 18-month founding term means the entire first cohort renews in the '
     'same month, on a code path that has never run against live traffic. Put '
     'it in a calendar now — it is both a revenue event and a support event, '
     'and it is the direct consequence of the founding term being longer.')

doc.save(OUT)
print('wrote', OUT)
