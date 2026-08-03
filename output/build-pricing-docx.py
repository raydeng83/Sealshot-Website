#!/usr/bin/env python3
"""Build the Sealshot v1.0 pricing & licensing reference as a .docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0xC2, 0x41, 0x0C)   # Sealshot orange
INK = RGBColor(0x0B, 0x12, 0x20)
MUTED = RGBColor(0x4F, 0x5A, 0x72)

doc = Document()

# ---------- base styles ----------
normal = doc.styles['Normal']
normal.font.name = 'Helvetica Neue'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before in [
    ('Heading 1', 20, INK, 22),
    ('Heading 2', 14, ACCENT, 16),
    ('Heading 3', 11.5, INK, 12),
]:
    s = doc.styles[name]
    s.font.name = 'Helvetica Neue'
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(5)
    s.paragraph_format.keep_with_next = True

for section in doc.sections:
    section.top_margin = section.bottom_margin = Inches(0.9)
    section.left_margin = section.right_margin = Inches(0.9)


def para(text='', style=None, bold=False, italic=False, color=None, size=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(*parts, style=None, space_after=None):
    """rich(('text', {'bold':True}), 'plain text', …)"""
    p = doc.add_paragraph(style=style)
    for part in parts:
        text, opts = part if isinstance(part, tuple) else (part, {})
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        r.italic = opts.get('italic', False)
        if 'color' in opts:
            r.font.color.rgb = opts['color']
        if opts.get('mono'):
            r.font.name = 'Menlo'
            r.font.size = Pt(9.5)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        p.add_run(bold_lead).bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def shade(cell, hex_fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def table(headers, rows, widths=None, right_align=()):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        shade(cell, 'F1F4F9')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if i in right_align:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Menlo'
        r.font.size = Pt(9)
        r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ======================= TITLE =======================
para('Sealshot', style='Heading 1', space_after=0)
para('Pricing, licensing, and renewals — v1.0', bold=True, color=ACCENT, size=13, space_after=2)
para('Reference for the v1.0 public release.', italic=True, color=MUTED, space_after=14)
para('Edition 2 · 3 August 2026', italic=True, color=MUTED, size=9.5, space_after=14)

# ======================= 1. AT A GLANCE =======================
para('1. At a glance', style='Heading 2')
para('Sealshot is sold as a perpetual license with a time-limited window for '
     'new releases. You buy once and keep the app forever. Updates published '
     'within your window are included; after it lapses you keep everything you '
     'have and may renew if you want newer releases.')
para('Every license includes every feature. There are no tiers, and nothing — '
     'Smart Redaction, encryption, recording, on-device AI — is withheld for a '
     'higher price.')

# ======================= 2. PRICING =======================
para('2. Pricing', style='Heading 2')
table(
    ['Offer', 'Price', 'What it covers'],
    [
        ['Individual license', '$49', 'One user, two Macs, commercial use, 12 months of updates'],
        ['Founding license', '$39', 'One user, two Macs, 18 months of updates. Available to buyers before the v1.0 release only'],
        ['Update renewal', '$24', 'Another 12 months of updates. Entirely optional'],
        ['Business volume', 'Contact us', 'Offline organization license, contracted seat count, one invoice'],
    ],
    widths=[1.5, 0.85, 4.3], right_align=(1,),
)
para('Prices are in USD, one-time unless stated. Tax is calculated at checkout.',
     italic=True, color=MUTED, size=9.5)

para('Volume pricing', style='Heading 3')
table(
    ['Users', 'Price per user'],
    [['1–9', '$49'], ['10–24', '$44'], ['25–99', '$39'], ['100 or more', 'Custom quote']],
    widths=[1.6, 1.6], right_align=(1,),
)

# ======================= 3. PERPETUAL vs UPDATES =======================
para('3. Perpetual access and the update window', style='Heading 2')
para('These are two separate things, and the distinction is the heart of the model.')

para('Perpetual app access', style='Heading 3')
para('Your license does not expire. Sealshot keeps working indefinitely, and your '
     'captures are always yours to view, edit, annotate, search, and export — '
     'regardless of license state. Your existing work is never withheld.')

para('The update window', style='Heading 3')
rich('Your license carries an ', ('Updates through', {'bold': True}),
     ' date. It grants you every Sealshot release published on or before that '
     'date. Installing a release published after it pauses ',
     ('new captures and recordings only', {'bold': True}),
     ' until you renew or return to a covered version.')
para('Every release remains downloadable, so a lapsed license can always '
     'reinstall the most recent version it covers.')

para('Terminology', style='Heading 3')
rich('Use ', ('“App access: Perpetual”', {'bold': True}), ' and ',
     ('“Updates through <date>”', {'bold': True}), '. Avoid “expiration date”, '
     '“valid until”, or “subscription” — the license never expires; only '
     'eligibility for newly published releases lapses.')

# ======================= 4. TRIAL =======================
para('4. Free trial', style='Heading 2')
bullet('', bold_lead='14 days, fully featured. ')
doc.paragraphs[-1].runs[-1].text = 'Every feature is available during the trial — nothing is held back.'
bullet('No account, no sign-up, no telemetry. Download and run.')
bullet('A reminder appears during the final week of the trial.')
rich('When the trial ends, ', ('new captures and recordings pause', {'bold': True}),
     '. Everything captured during the trial stays fully viewable, editable, and '
     'exportable. Activate a license at any time to resume.')

# ======================= 5. WHAT YOU RECEIVE =======================
para('5. What you receive when you buy', style='Heading 2')
para('Two emails arrive, from two senders:')
table(
    ['Email', 'From', 'Contains'],
    [
        ['Payment receipt', 'Polar (our merchant of record)', 'Your invoice, including any tax collected'],
        ['Your license', 'Sealshot', 'The .sealshotlicense file, attached, with activation instructions'],
    ],
    widths=[1.35, 2.15, 3.15],
)
para('Keep the license file. It is your proof of purchase, it is personally '
     'identifying, and it is the file you use to activate on each of your Macs.')

# ======================= 6. THE LICENSE FILE =======================
para('6. The license file', style='Heading 2')
para('A license is a readable text file ending in a single signed line. The '
     'readable part is cryptographically bound to the signature, so editing any '
     'of it — including deleting the closing paragraph — invalidates the license.')

para('Individual license', style='Heading 3')
code_block([
    'Sealshot License',
    '================',
    'Licensed to:      Jane Doe',
    'Email:            jane@example.com',
    'License ID:       5B4894D2-7DD5-4BDA-89D2-DC30954E2EE4',
    'License type:     Individual',
    'License issued:   2027-08-01',
    'App access:       Perpetual',
    'Updates through:  2028-09-15',
    'Users:            1',
    'Macs per user:    2',
    '',
    'This license does not expire. It permits use of every Sealshot',
    'release whose entitlement date is on or before 2028-09-15.',
    '',
    'Keep this file exactly as received. The information above is',
    'cryptographically signed; modifying it invalidates the license.',
    '',
    'SEALSHOT1.<signed-license-data>',
])

para('Business volume license', style='Heading 3')
code_block([
    'Sealshot License',
    '================',
    'Licensed to:      Acme, Inc.',
    'Purchaser email:  software@acme.example',
    'License ID:       946C87B1-C2D7-46E9-B8CF-D88EB86532EE',
    'License type:     Business Volume',
    'License issued:   2027-08-01',
    'App access:       Perpetual',
    'Updates through:  2028-09-15',
    'User seats:       10',
    'Macs per user:    2',
    '',
    'This is an offline, organization-wide license for up to 10 users.',
    'Sealshot does not transmit installation or usage information.',
    '',
    'SEALSHOT1.<signed-license-data>',
])

# ======================= 7. ACTIVATION =======================
para('7. Activation', style='Heading 2')
for i, step in enumerate([
    'Open Settings ▸ License in Sealshot.',
    'Open the license file, or drag it onto the window.',
    'Sealshot verifies the signature entirely on your Mac.',
    'The license is stored, and the License screen shows your update date.',
], 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    p.add_run(step)
rich('Activation is ', ('offline', {'bold': True}),
     '. There is no activation server, no device fingerprinting, and no '
     'registration. Sealshot never reports that you installed it.')
para('Use the same file on both of your Macs.')

# ======================= 8. RENEWAL =======================
para('8. Renewing your update window', style='Heading 2')
para('Renewal is optional. Everything you own keeps working without it.')
rich('A renewal is a ', ('replacement license file', {'bold': True}),
     ', not a second license. It carries the same license ID, the same owner and '
     'seat count, a new issue date, and an extended ',
     ('Updates through', {'bold': True}), ' date.')

para('Renewing early keeps your unused time', style='Heading 3')
para('Renewing before your window ends adds twelve months to the date you '
     'already have — you never lose time by renewing early.')
table(
    ['Situation', 'Current window ends', 'You renew on', 'New window ends'],
    [
        ['Early renewal', '2027-09-15', '2027-08-01', '2028-09-15'],
        ['After it lapsed', '2027-09-15', '2028-03-01', '2029-03-01'],
    ],
    widths=[1.5, 1.75, 1.5, 1.75],
)

para('Installing a renewal', style='Heading 3')
para('Open the renewal file the same way you activated the original. Sealshot '
     'replaces the stored license and shows the new date immediately. Sealshot '
     'protects you from going backwards:')
table(
    ['File you open', 'What happens'],
    [
        ['A newer renewal for this license', 'Installed, and the new date takes effect'],
        ['The same file again', '“This license is already installed.”'],
        ['An older file for this license', 'Rejected — “A newer renewal license is already installed.”'],
        ['A license with a different ID', 'Sealshot asks before replacing your current license'],
    ],
    widths=[2.6, 4.05],
)
para('So re-opening an old attachment can never shorten your coverage.',
     italic=True, color=MUTED, size=9.5)

# ======================= 9. BUSINESS =======================
para('9. Refunds', style='Heading 2')

para('Thirty days, no reason required. Email support@seal-shot.com within 30 '
     'days of buying and the purchase is refunded in full. Renewals are covered '
     'by the same 30 days.')

para('Purchases are processed by Polar as merchant of record, so the refund '
     'returns through them to the card used. A reason is welcome but never a '
     'condition.')

para('A refund revokes the license, and Sealshot returns to its unlicensed '
     'state. Nothing captured is deleted or locked — the library stays on the '
     'customer\u2019s Mac, readable and exportable, exactly as when a trial ends. '
     'That distinction is worth stating plainly, because it is the question a '
     'buyer actually has and most software answers it badly.')

para('Volume licenses are quoted and invoiced individually, so their terms are '
     'agreed in the quote rather than set here.')

para('10. Business volume licensing', style='Heading 2')
para('For organizations, Sealshot is licensed per named user under a contract, '
     'with no activation infrastructure — consistent with a product that '
     'transmits nothing.')
bullet('one named user, on up to two Macs.', bold_lead='A seat is ')
bullet('containing the organization name, seat count, and update date.',
       bold_lead='You receive one signed license file ')
bullet('internally; there is nothing to deploy per machine.',
       bold_lead='Distribute that file ')
bullet('when people join or leave.', bold_lead='Seats are reassignable ')
bullet('no activation reporting, no device tracking, no usage data.',
       bold_lead='Sealshot performs ')
bullet('rather than technically enforced.', bold_lead='Compliance is contractual ')
rich(('Seat changes are handled by reissue. ', {'bold': True}),
     'Adding seats produces a new license file with the higher count; the '
     'previous file continues to work and continues to show the old count, so '
     'replace it when you distribute the new one.')
para('To arrange a volume license, contact us with your seat count and billing '
     'details. We will provide a quote and an invoice.')

# ======================= 10. PRIVACY =======================
para('11. Privacy of the licensing system', style='Heading 2')
para('Licensing is built to the same standard as the rest of Sealshot:')
bullet('No account is ever created.')
bullet('Verification happens on your Mac. Your license file is never uploaded.')
bullet('No activation, installation, or usage information is transmitted.')
bullet('The app’s only network activity is checking for updates and, with your '
       'explicit consent, downloading the optional on-device redaction model.')
para('The trade-off is stated plainly: because nothing is tracked, honoring '
     'seat counts is a matter of trust rather than enforcement.')

# ======================= 11. FAQ =======================
para('12. Common questions', style='Heading 2')
faq = [
    ('Does Sealshot stop working when my update window ends?',
     'No. The app and all your captures keep working indefinitely. Only new '
     'captures on a release published after your window are paused — and '
     'reinstalling your last covered version restores those too.'),
    ('Is this a subscription?',
     'No. You buy the license once. Renewal buys another year of new releases '
     'and is entirely optional.'),
    ('How many Macs can I use?',
     'Two, for your own use. Install the same license file on both.'),
    ('Can I use Sealshot commercially?',
     'Yes. The individual license includes commercial use.'),
    ('What if I lose my license file?',
     'Contact us with your purchase details and we will reissue it.'),
    ('Do I get v2 for free?',
     'If it is published inside your update window, yes. There is no separate '
     'major-version upgrade charge.'),
    ('Which features do I get?',
     'All of them, on every license — including Smart Redaction, Enhanced '
     'Security encryption, screen recording, and the on-device AI features '
     'your Mac supports.'),
]
for q, a in faq:
    para(q, style='Heading 3')
    para(a)

out = '/Users/ledeng/projects/Sealshot-Website/output/Sealshot-Pricing-and-Licensing-v1.0.docx'
doc.save(out)
print('wrote', out)
