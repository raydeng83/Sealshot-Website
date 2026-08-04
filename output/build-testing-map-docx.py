#!/usr/bin/env python3
"""
Build output/Sealshot-Testing-Map.docx from docs/testing-map.md.

Parses the markdown rather than restating it, so the Word doc cannot drift from
the source the way a hand-authored copy would — regenerate after any edit and the
two agree by construction. The other generators in this folder hand-author their
content; this one does not, because testing-map.md is itself the deliverable and
is still being revised.

Screenshots are inserted after specific scenario cards, keyed by heading text in
SHOTS below. A missing file is reported and skipped rather than failing the build,
so the doc still generates on a machine without the full public/manual set.

    python3 output/build-testing-map-docx.py
"""
import pathlib
import re
from datetime import date

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
SRC = ROOT / 'docs' / 'testing-map.md'
SHOT_DIR = ROOT / 'public' / 'manual'
OUT = pathlib.Path(__file__).resolve().parent / 'Sealshot-Testing-Map.docx'

ACCENT = RGBColor(0xC2, 0x41, 0x0C)
INK = RGBColor(0x0B, 0x12, 0x20)
MUTED = RGBColor(0x4F, 0x5A, 0x72)

# Heading text -> (image, caption). Cards with nothing worth showing are absent:
# scenario 16 is a packet capture in another app, and 17 is a room full of Macs.
SHOTS = {
    'The app': ('editor-overview.jpg', 'The editor — capture open, annotation toolbar, Info panel.'),
    'Capture, Annotate and Export': ('capture-area.jpg', 'Area capture, with the pixel loupe and live dimensions.'),
    'Scroll Capture and Verify': ('capture-scrolling.jpg', 'Scrolling capture. Check the seams at 100%, not at thumbnail size.'),
    'Capture Straight to File': ('capture-destination.png', 'Where captures land, set once in Settings.'),
    'Record with Narration, Pause and Resume': ('record-prompt.jpg', 'The record prompt — audio sources, cursor, countdown.'),
    'Extract from a Recording': ('extract-data.jpg', 'Structured extraction pulling a table out of pixels.'),
    'Redact, Export and Attack': ('redaction-review.png', 'Smart Redaction review. Detected items arrive pre-checked.'),
    'Work with Encryption On': ('lock-screen.png', 'Locked. Captures keep working; viewing waits for Touch ID.'),
    'Lose the Code and Recover': ('security-recovery-key.png', 'The recovery code ceremony — the one screen worth photographing.'),
    'Share with One Recipient': ('export-package.png', 'Export as a package, with passcode and expiry.'),
    'Build a Library and Search It': ('library-search.jpg', 'Full-text search over the text inside the images.'),
    'Delete and Recover': ('library-deleted.jpg', 'Recently Deleted, before the timed purge.'),
}

doc = Document()

normal = doc.styles['Normal']
normal.font.name = 'Helvetica Neue'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before in [
    ('Heading 1', 20, INK, 22),
    ('Heading 2', 15, ACCENT, 20),
    ('Heading 3', 12, INK, 14),
    ('Heading 4', 11, INK, 13),
]:
    st = doc.styles[name]
    st.font.name = 'Helvetica Neue'
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True

for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    sec.left_margin = sec.right_margin = Inches(0.9)

CONTENT_IN = 6.7  # letter width less both margins


# ── inline formatting ────────────────────────────────────────────────────────
TOKEN = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)')


def add_inline(p, text):
    """Render **bold**, `code`, [links](url) and *italic* into a paragraph."""
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Menlo'
            r.font.size = Pt(9.5)
        elif part.startswith('[') and '](' in part:
            label = part[1:part.index('](')]
            r = p.add_run(label)
            r.font.color.rgb = ACCENT
        elif part.startswith('*') and part.endswith('*'):
            p.add_run(part[1:-1]).italic = True
        else:
            p.add_run(part)


def add_image(name, caption):
    path = SHOT_DIR / name
    if not path.exists():
        print(f'  ! missing screenshot, skipped: {name}')
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(4.9))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    return True


def set_orientation(landscape):
    """New section, rotated, for tables too wide for portrait.

    UNEXERCISED as of this commit: the coverage matrix that needed it (13
    columns) has since been removed from the source, so nothing currently
    triggers this path. Kept because a wide table is likely to come back, but
    treat it as untested until one does."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.top_margin = sec.bottom_margin = Inches(0.7)
    sec.left_margin = sec.right_margin = Inches(0.7)
    w, h = sec.page_width, sec.page_height
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        if w < h:
            sec.page_width, sec.page_height = h, w
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        if w > h:
            sec.page_width, sec.page_height = h, w
    return sec


def add_table(rows):
    header, body = rows[0], rows[2:]          # rows[1] is the |---| separator
    wide = len(header) > 8
    if wide:
        set_orientation(True)
    t = doc.add_table(rows=1, cols=len(header))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    size = Pt(7) if wide else Pt(9.5)
    for i, cell in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(cell)
        r.bold = True
        r.font.size = size
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
    for row in body:
        cells = t.add_row().cells
        for i, val in enumerate(row[:len(header)]):
            cells[i].text = ''
            add_inline(cells[i].paragraphs[0], val)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = size
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
    if wide:
        set_orientation(False)


# ── parse ────────────────────────────────────────────────────────────────────
lines = SRC.read_text().split('\n')

# title page
doc.add_paragraph(lines[0].lstrip('# ').strip(), style='Heading 1')
sub = doc.add_paragraph()
r = sub.add_run('Scenario testing brief · Sealshot 0.7.3 (17)')
r.font.size = Pt(11.5)
r.font.color.rgb = MUTED
stamp = doc.add_paragraph()
r = stamp.add_run(f'Generated {date.today():%B %-d, %Y} from docs/testing-map.md')
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = MUTED

i, shots_used, pending_shot = 1, 0, None
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # a screenshot goes after the card's whole body, not between title and steps
    if stripped.startswith('#') or (pending_shot and stripped == '' and i + 1 < len(lines)
                                    and lines[i + 1].startswith('#')):
        if pending_shot and stripped.startswith('#'):
            if add_image(*pending_shot):
                shots_used += 1
            pending_shot = None

    if not stripped:
        i += 1
        continue

    if stripped == '---':
        i += 1
        continue

    m = re.match(r'^(#{2,4}) (.+)$', stripped)
    if m:
        level, text = len(m.group(1)), m.group(2)
        doc.add_paragraph(text, style=f'Heading {level}')
        if text in SHOTS:
            pending_shot = SHOTS[text]
        i += 1
        continue

    if stripped.startswith('|'):
        rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
            i += 1
        add_table(rows)
        continue

    m = re.match(r'^(\d+)\. (.+)$', stripped)
    if m or stripped.startswith('- '):
        style = 'List Number' if m else 'List Bullet'
        text = m.group(2) if m else stripped[2:]
        # fold continuation lines (indented, no marker) into the same item
        i += 1
        while (i < len(lines) and lines[i].startswith('   ')
               and not re.match(r'^\s*(\d+\.|-) ', lines[i])):
            text += ' ' + lines[i].strip()
            i += 1
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        add_inline(p, text)
        continue

    # plain paragraph, folding its wrapped lines back together
    text = stripped
    i += 1
    while (i < len(lines) and lines[i].strip()
           and not lines[i].strip().startswith(('#', '-', '|', '---'))
           and not re.match(r'^\d+\. ', lines[i].strip())):
        text += ' ' + lines[i].strip()
        i += 1
    p = doc.add_paragraph()
    # the *~10 min · workbooks …* metadata line
    if text.startswith('*') and text.endswith('*') and text.count('*') == 2:
        r = p.add_run(text[1:-1])
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        p.paragraph_format.space_after = Pt(4)
    else:
        add_inline(p, text)

if pending_shot and add_image(*pending_shot):
    shots_used += 1

doc.save(OUT)
kb = OUT.stat().st_size / 1024
print(f'wrote {OUT.relative_to(ROOT)}  ({kb:.0f} KB, {shots_used} screenshots)')
