#!/usr/bin/env python3
"""
Build output/Sealshot-Support-Email-Templates.docx.

Canned replies for support@seal-shot.com, in the same voice as the purchase
email the customer already received. US English throughout.

    python3 output/build-support-templates-docx.py
"""
import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0xC2, 0x41, 0x0C)
INK = RGBColor(0x0B, 0x12, 0x20)
MUTED = RGBColor(0x4F, 0x5A, 0x72)

OUT = pathlib.Path(__file__).resolve().parent / 'Sealshot-Support-Email-Templates.docx'

doc = Document()

normal = doc.styles['Normal']
normal.font.name = 'Helvetica Neue'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before in [
    ('Heading 1', 20, INK, 22), ('Heading 2', 14, ACCENT, 18), ('Heading 3', 11.5, INK, 14),
]:
    s = doc.styles[name]
    s.font.name = 'Helvetica Neue'
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.keep_with_next = True

for section in doc.sections:
    section.top_margin = section.bottom_margin = Inches(0.9)
    section.left_margin = section.right_margin = Inches(0.9)


def para(text='', style=None, bold=False, italic=False, color=None, size=None, after=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold, r.italic = bold, italic
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = Pt(size)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def rich(*parts, after=None):
    p = doc.add_paragraph()
    for part in parts:
        text, opts = part if isinstance(part, tuple) else (part, {})
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        r.italic = opts.get('italic', False)
        if 'color' in opts:
            r.font.color.rgb = opts['color']
        if opts.get('mono'):
            r.font.name = 'Menlo'
            r.font.size = Pt(9.5)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def bullet(*parts):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    for part in parts:
        text, opts = part if isinstance(part, tuple) else (part, {})
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        if opts.get('mono'):
            r.font.name = 'Menlo'
            r.font.size = Pt(9.5)
    return p


def shade(cell, fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(el)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        shade(c, 'F1F4F9')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def template(title, when, subject, body, note=None):
    """One canned reply: heading, when to send, subject line, body, optional note."""
    para(title, style='Heading 3')
    para(when, italic=True, color=MUTED, size=9.5, after=6)

    rich(('Subject:  ', {'bold': True}), (subject, {'mono': True}), after=4)

    # The body sits in a single shaded cell so it reads as a block to paste.
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    shade(cell, 'F7FAFC')
    cell.width = Inches(6.5)
    cell.text = ''
    for i, block in enumerate(body):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(block)
        r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    if note:
        rich(('Note. ', {'bold': True, 'color': ACCENT}), (note, {}), after=10)


# ═══════════════════════════════════════════════════════════════════ title ═══
para('Sealshot', style='Heading 1', after=0)
para('Support email templates', bold=True, color=ACCENT, size=13, after=2)
para('For support@seal-shot.com · Edition 1, 3 August 2026',
     italic=True, color=MUTED, after=14)

para('Canned replies in the same voice as the purchase email a customer has '
     'already received. Adapt freely — a template that gets sent unchanged '
     'every time starts to read like a form letter, which is the thing these '
     'are meant to avoid.')

# ══════════════════════════════════════════════════════════════════ 1 rules ══
para('1. Rules', style='Heading 2')

para('Five that matter more than the wording of any individual reply.')

para('Never ask for a screenshot without saying "throwaway content"', style='Heading 3')
para('This is a screenshot app. Any reproduction a customer sends will contain '
     'their real work. Asking for one carelessly contradicts the product\'s '
     'entire premise; asking for one carefully demonstrates it.')

para('No open or click tracking, ever', style='Heading 3')
para('Gmail adds none. But most helpdesk products enable pixel tracking by '
     'default, so this becomes a real decision the day you adopt one. A '
     'privacy-first product measuring whether you opened its support reply is '
     'the most quotable contradiction available to a critic.')

para('Lead with the self-serve fix, then offer to do it', style='Heading 3')
para('Most license problems are "check your spam folder" or "use the original '
     'attachment". Giving the customer the fast path first respects their time '
     'and cuts your own volume.')

para('No apology theater, no ticket numbers', style='Heading 3')
para('One person answering quickly is an advantage over the companies these '
     'customers are used to. Don\'t disguise it as a queue. Say what you will '
     'do and when; never "we\'ll look into it".')

para('Be the one who says the unwelcome thing plainly', style='Heading 3')
para('Two of the replies below deliver bad news — an unrecoverable library, an '
     'update window that has closed. Under pressure the instinct is to soften '
     'into "let me see what I can do". False hope makes both situations worse, '
     'and a customer told the truth immediately often trusts the product more '
     'afterwards, not less.')

# ═════════════════════════════════════════════════════════════ 2 signature ═══
para('2. Signature and placeholders', style='Heading 2')

para('Keep the signature to two lines. No legal footer, no social icons, no '
     '"sent from my…".')

rich(('— Ray', {'mono': True}), after=2)
rich(('Sealshot · seal-shot.com/docs', {'mono': True}), after=8)

para('The purchase email signs off "— Sealshot". Either match that or use your '
     'name consistently — but don\'t alternate, or a customer with two emails '
     'from you will wonder how many people there are.')

table(
    ['Placeholder', 'Means'],
    [['{name}', "Customer's first name; omit the line entirely rather than write “Hi there” twice in a thread"],
     ['{version}', 'Newest release their license covers, e.g. 1.2.0'],
     ['{date}', 'Their Updates through date, written out: 2 February 2028'],
     ['{id}', 'License ID — only when you need them to confirm it']],
    widths=[1.2, 5.3])

# ═════════════════════════════════════════════════════════════ 3 templates ═══
para('3. Templates', style='Heading 2')

template(
    'License didn\'t arrive',
    'The most common and most urgent — they have paid and received nothing. '
    'Look the order up before replying, and attach the reissued file so the '
    'reply resolves it in one round trip.',
    'Your Sealshot license',
    ["Hi {name},",
     "Sorry — that should have arrived within a minute of your purchase.",
     "First worth checking: search your mail for “Sealshot license”, including "
     "spam. It arrives as a file attachment from a new sender, which is exactly "
     "what spam filters dislike.",
     "If it isn't there, your license is attached to this email. To activate it, "
     "open Sealshot, go to Settings ▸ License, and open the attached file or "
     "drag it onto the window.",
     "Keep the file — it's your proof of purchase, and it's what you'll use on a "
     "new Mac.",
     "— Ray"],
    note='Attach the file before sending. A reply that only says "check your spam" '
         'and makes them write back doubles the wait on the one issue where the '
         'customer is already out of pocket.')

template(
    'License won\'t activate',
    'They see "That doesn\'t look like a valid license file." Nearly always the '
    'wrong file, or one that has been opened and re-saved.',
    'Re: Sealshot license',
    ["Hi {name},",
     "That message usually means the file has been altered slightly — often by "
     "being opened in a text editor and saved again, which can change invisible "
     "characters even if the text looks identical.",
     "The fix is to use the original attachment from your purchase email, "
     "downloaded fresh rather than copied out. The file ends in "
     ".sealshotlicense.",
     "If that still doesn't work, tell me and I'll reissue it — no need to send "
     "the file back.",
     "— Ray"],
    note='Don\'t ask them to send a screenshot of the error. The window may show '
         'their name, email and library behind it. The exact wording is enough, '
         'and there are only three possible messages.')

template(
    'License file was modified',
    'They see "This license file has been modified and is no longer valid." A '
    'different, more specific failure than the one above.',
    'Re: Sealshot license',
    ["Hi {name},",
     "That one is specific: the readable text at the top of the license file has "
     "changed since it was issued. The details are cryptographically signed, so "
     "editing even a space breaks the signature — which is what stops anyone "
     "editing an expiry date or a name.",
     "Use the original attachment from your purchase email and it will activate. "
     "If you no longer have it, say so and I'll send a fresh one.",
     "— Ray"])

template(
    'License revoked',
    'They see "This license has been revoked." Rare, and you should already '
    'know why. Do not send this without checking the reason first.',
    'Re: Sealshot license',
    ["Hi {name},",
     "That license has been revoked, which is why it won't activate.",
     "If that's unexpected, tell me and I'll look into it today — revocation is "
     "deliberate and rare, so if it has happened by mistake I want to know.",
     "— Ray"],
    note='Never send this as a first reply to a confused customer. Establish why '
         'the license was revoked, then write specifically. A generic revocation '
         'notice to someone who did nothing wrong is worse than a slow reply.')

template(
    'Update window has ended',
    'They see that a newer release is outside their update window. The delicate '
    'one — a customer may believe something has been taken away, so the first '
    'line has to remove that fear before anything is explained.',
    'Your Sealshot license and updates',
    ["Hi {name},",
     "Sealshot hasn't stopped working and won't — your license is perpetual. "
     "What's ended is the window for new versions, so {version} is the newest "
     "release your license covers, and it keeps running indefinitely.",
     "If you'd like the newer releases, renewing is $24 for another 12 months: "
     "seal-shot.com/renew. Renewing early doesn't lose you anything — unused "
     "time is added on top of what you have.",
     "And if you'd rather not, that's genuinely fine. Nothing you've captured "
     "becomes inaccessible, and there's no reduced mode to work around.",
     "— Ray"],
    note='The last paragraph is the whole point of the pricing model. Say it '
         'plainly or customers will not believe it — most software has trained '
         'them to expect a lockout.')

template(
    'New Mac, or reinstalling',
    'Moving machines, or a clean install.',
    'Re: Sealshot on a new Mac',
    ["Hi {name},",
     "Use the same license file — it isn't tied to a particular Mac. Download "
     "Sealshot from seal-shot.com/download, then open the license file or drag "
     "it onto Settings ▸ License.",
     "One license covers one person on up to two Macs, so you don't need to "
     "deactivate the old one first.",
     "If you no longer have the file, tell me the email you bought with and "
     "I'll resend it.",
     "— Ray"])

template(
    'Refund request',
    'Polar is the merchant of record, so the transaction is theirs.',
    'Re: Sealshot refund',
    ["Hi {name},",
     "Happy to sort that out. Purchases are handled by Polar, our payment "
     "provider, so the refund goes through them — reply to your Polar receipt, "
     "or tell me the email you bought with and I'll start it from this side.",
     "If something specific didn't work, I'd genuinely like to hear it, whether "
     "or not you want the refund either way.",
     "— Ray"],
    note='Ask what went wrong, but never make answering a condition of the '
         'refund. Refunds are cheap; a public complaint about being '
         'interrogated for one is not.')

template(
    'Encrypted library, recovery key lost',
    'The worst case. This template exists to stop the truth being softened '
    'under pressure.',
    'Re: Sealshot encrypted library',
    ["Hi {name},",
     "I have to be straight with you: without the recovery key, that library "
     "can't be opened. Not by me either — the encryption keys never leave your "
     "Mac, and we hold nothing that could unlock it. That's the design, and it's "
     "why the warnings when you turned it on were as blunt as they were.",
     "I'm sorry. If any captures were exported before encryption was enabled, "
     "those are unaffected.",
     "— Ray"],
    note='Do not write "let me check what I can do." There is nothing to check, '
         'and a day of false hope makes the same answer land far harder.')

template(
    'Bug report received',
    'Acknowledge within a day, even if you have not diagnosed it. Silence is '
    'what makes people give up on reporting.',
    'Re: {their subject}',
    ["Hi {name},",
     "Thanks — that's a real one, and I've reproduced it / I'll try to reproduce "
     "it today.",
     "If it's straightforward it'll be in the next release, and I'll email you "
     "when it ships. If it turns out to be more involved, I'll tell you that "
     "rather than leave you guessing.",
     "— Ray"],
    note='Pick one of the two clauses in the second line before sending. And if '
         'you promise to email when it ships, keep a note — an unkept promise '
         'here costs more than never having made it.')

template(
    'Crash report needed',
    'When a crash is reported without a report attached.',
    'Re: Sealshot crashed',
    ["Hi {name},",
     "Could you send the crash report? macOS writes one locally and it never "
     "leaves your Mac unless you send it.",
     "In Finder, choose Go ▸ Go to Folder… and enter "
     "~/Library/Logs/DiagnosticReports, then look for the newest file starting "
     "with Sealshot. Attach that .ips file to a reply.",
     "It contains stack traces, the app version and your macOS version — no "
     "screenshots and nothing you captured. You're welcome to open it in a text "
     "editor and read it before sending.",
     "— Ray"],
    note='The last paragraph is the part that gets the file sent. Users of a '
         'privacy-first app are rightly suspicious of "send us your logs", and '
         'inviting them to read it first answers that in one sentence.')

template(
    'Feature request',
    'Answer honestly: yes, no, or not soon. Never "we\'ll consider it".',
    'Re: {their subject}',
    ["Hi {name},",
     "Thanks for this — noted properly, not filed away.",
     "[Yes: it's on the list for a coming release, though I won't guess at a "
     "date.] [Not soon: it's a reasonable idea, but it isn't close to the top, "
     "and I'd rather say so than leave you waiting.] [No: it would work against "
     "how Sealshot handles X, so I don't expect to build it — here's why: …]",
     "— Ray"],
    note='Choose one bracket and delete the rest. A clear no is more respected '
         'than a vague maybe, and it costs you nothing — the person asked '
         'because they care about the product.')

template(
    'Volume or business licensing',
    'A business asking about multiple seats. Quoted and invoiced by hand; there '
    'is no volume product at checkout.',
    'Sealshot volume licensing',
    ["Hi {name},",
     "Yes — volume licenses are handled directly, not through the website.",
     "Pricing is per user: $44 each for 10–24, $39 each for 25–99, and "
     "custom above 100. Each license covers one person on up to two Macs and "
     "includes 12 months of updates, the same as an individual license.",
     "It's issued as a single organization-wide license file, so there's nothing "
     "to activate per machine and no license server to run. Tell me the seat "
     "count and how you'd like to be invoiced, and I'll send a quote.",
     "— Ray"],
    note='These per-seat figures are NOT published on the site — the FAQ says '
         '"volume discounts as the count grows" and nothing more, so this reply '
         'is where a prospect first sees a number. Check them against the '
         'pricing document before quoting; if that changes, this template is '
         'the thing most likely to be forgotten. Also worth saying out loud: '
         'Sealshot has no license server and no per-seat activation, which to a '
         'company that has managed software licensing before is a selling point '
         'rather than a limitation.')

template(
    'Do you send my screenshots anywhere?',
    'The best brand moment you get. Answer with specifics, not reassurance.',
    'Re: privacy',
    ["Hi {name},",
     "No. Capture, recording, OCR, redaction, the AI features and search all run "
     "on your Mac. There's no account, no telemetry, and nothing you capture is "
     "uploaded.",
     "Sealshot makes exactly two network requests, both downloads and neither "
     "carrying your content: a daily check for updates, and — only if you turn "
     "it on — a one-time download of the enhanced redaction model. You can turn "
     "the update check off in Settings.",
     "The full detail is at seal-shot.com/privacy, including what we store when "
     "you buy a license, which is your name, email and license dates.",
     "— Ray"],
    note='Mentioning what you DO store is what makes the rest credible. A reply '
         'that only says "we store nothing" reads like marketing; one that names '
         'the exceptions reads like an answer.')

template(
    'Trial expired, how do I buy',
    'Straightforward, and worth answering fast.',
    'Re: Sealshot trial',
    ["Hi {name},",
     "seal-shot.com/buy — it's a one-time purchase, and the license arrives by "
     "email as a file within a minute. Open it or drag it onto Settings ▸ "
     "License and you're done.",
     "Everything you captured during the trial stays where it is and remains "
     "accessible; activating doesn't migrate or re-import anything.",
     "— Ray"],
    note='The second paragraph pre-empts the real worry: that trial work is '
         'trapped or will be lost. Nobody asks it outright.')

# ════════════════════════════════════════════════════════════════ 4 storing ══
para('4. Where to keep these', style='Heading 2')

para('Gmail ▸ Settings ▸ Advanced ▸ enable Templates. Then compose the '
     'reply, and from the ⋮ menu choose Save draft as template.')

rich(('Templates are per-account, which interacts with how you set '
      'support@ up. ', {'bold': True}),
     ('If it is a real mailbox with delegation, both people share one set — '
      'edit once and both benefit. If it is a Google Group, each member keeps '
      'their own copy and they drift apart within a month; in that case keep '
      'this document as the source and re-import after any edit.', {}))

para('5. Never include', style='Heading 2')
bullet('Ticket or case numbers. There is no queue; inventing one is theater.')
bullet('Tracking pixels or click-wrapped links.')
bullet('"As per our policy" — you are the policy.')
bullet('A request for a screenshot, unless you have said throwaway content.')
bullet('A promise to follow up that you have not written down somewhere.')
bullet('Two different sign-offs across one thread.')

doc.save(OUT)
print('wrote', OUT)
