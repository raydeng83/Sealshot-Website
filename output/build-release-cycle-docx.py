#!/usr/bin/env python3
"""Build the Sealshot v1.0 release-cycle reference as a .docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0xC2, 0x41, 0x0C)
INK = RGBColor(0x0B, 0x12, 0x20)
MUTED = RGBColor(0x4F, 0x5A, 0x72)

doc = Document()

normal = doc.styles['Normal']
normal.font.name = 'Helvetica Neue'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before in [
    ('Heading 1', 20, INK, 22),
    ('Heading 2', 14, ACCENT, 16),
    ('Heading 3', 11.5, INK, 12),
]:
    s = doc.styles[name]
    s.font.name = 'Helvetica Neue'
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(5)
    s.paragraph_format.keep_with_next = True

for section in doc.sections:
    section.top_margin = section.bottom_margin = Inches(0.9)
    section.left_margin = section.right_margin = Inches(0.9)


def para(text='', style=None, bold=False, italic=False, color=None, size=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = color
        if size is not None:
            r.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(*parts, style=None, space_after=None):
    p = doc.add_paragraph(style=style)
    for part in parts:
        text, opts = part if isinstance(part, tuple) else (part, {})
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        r.italic = opts.get('italic', False)
        if 'color' in opts:
            r.font.color.rgb = opts['color']
        if opts.get('mono'):
            r.font.name = 'Menlo'
            r.font.size = Pt(9.5)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(*parts, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(3)
    for part in parts:
        text, opts = part if isinstance(part, tuple) else (part, {})
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        if opts.get('mono'):
            r.font.name = 'Menlo'
            r.font.size = Pt(9.5)
    return p


def numbered(*parts):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    for part in parts:
        text, opts = part if isinstance(part, tuple) else (part, {})
        r = p.add_run(text)
        r.bold = opts.get('bold', False)
        if opts.get('mono'):
            r.font.name = 'Menlo'
            r.font.size = Pt(9.5)
    return p


def shade(cell, hex_fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def table(headers, rows, widths=None, mono_cols=()):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        shade(cell, 'F1F4F9')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if i in mono_cols:
                r.font.name = 'Menlo'
                r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Menlo'
        r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ======================= TITLE =======================
para('Sealshot', style='Heading 1', space_after=0)
para('The release cycle', bold=True, color=ACCENT, size=13, space_after=2)
para('How a Sealshot release is versioned, built, published, and delivered — v1.0 onward.',
     italic=True, color=MUTED, space_after=2)
para('Edition 2 · 3 August 2026', italic=True, color=MUTED, size=9.5, space_after=14)

# ======================= 1. OVERVIEW =======================
para('1. Overview', style='Heading 2')
para('Sealshot ships as a signed, notarized direct download that updates itself. '
     'One command builds, signs, notarizes, packages, publishes, and announces a '
     'release; the same command refuses to run if the release notes are missing.')
para('Three things are produced by every release, and they must stay consistent '
     'with each other:')
bullet(('A notarized .dmg', {'bold': True}), ' — what a new user downloads.')
bullet(('A signed update .zip plus an appcast entry', {'bold': True}),
       ' — how existing installs update themselves.')
bullet(('Release notes', {'bold': True}),
       ' — shown in the update dialog, on the GitHub release, and on the website changelog.')

# ======================= 2. VERSIONING =======================
para('2. Versioning', style='Heading 2')
para('Sealshot uses three-part versions. Every release also carries a build '
     'number that increases monotonically and never resets.')
table(
    ['Component', 'Meaning', 'Example'],
    [
        ['Major', 'A significant release. Included free if published inside your update window — there is no separate upgrade charge.', '1.0.0 → 2.0.0'],
        ['Minor', 'New features and capabilities.', '1.0.0 → 1.1.0'],
        ['Patch', 'Fixes and refinements, no new capability.', '1.1.0 → 1.1.1'],
        ['Build', 'Increments every published release. Used by the updater to compare versions.', '14 → 15'],
    ],
    widths=[0.85, 4.6, 1.2],
)
rich('The marketing version is held in ', ('app/project.yml', {'mono': True}),
     ' as ', ('MARKETING_VERSION', {'mono': True}),
     ' and is read automatically at release time. The build number is what '
     'Sparkle compares, so it must increase even when the marketing version '
     'does not.')

# ======================= 3. ENTITLEMENT DATE =======================
para('3. The entitlement date', style='Heading 2')
para('This is the mechanism that connects releases to licensing, and it is worth '
     'understanding before anything else.')
rich('Every published build is stamped at build time with the UTC day it was '
     'released, written into ', ('Info.plist', {'mono': True}), ' as ',
     ('SealshotReleaseDate', {'mono': True}), '. That is the release’s ',
     ('entitlement date', {'bold': True}), '.')
para('A license covers every release whose entitlement date falls on or before '
     'its “Updates through” date. That comparison is the whole of the update '
     'window: no server is consulted, and no clock other than the two dates '
     'matters.')
table(
    ['Build type', 'Entitlement date', 'Effect'],
    [
        ['Published direct release', 'Stamped with the release day', 'Compared against the license’s update window'],
        ['Development build', 'Empty', 'Licensing fails open — never gates a dev build'],
    ],
    widths=[1.6, 1.75, 3.3],
)
para('Practical consequence: the date is fixed when the release is built, not '
     'when a user installs it. A customer who installs a two-year-old version '
     'today is still covered if that version’s entitlement date sits inside '
     'their window.', italic=True, color=MUTED, size=9.5)

# ======================= 4. RELEASE NOTES =======================
para('4. Release notes are written continuously', style='Heading 2')
rich('Notes are not written at release time. Any change that alters '
     'user-visible behavior adds a bullet to ',
     ('docs/release-notes/next.md', {'mono': True}), ' as it is made.')

para('How to write them', style='Heading 3')
bullet('Describe what changed for the user, not how it was implemented.')
bullet('Group bullets under ## headings — add or drop headings as the release needs.')
bullet('Name the visible thing: a menu item, a shortcut, a setting.')
bullet('Fixes describe the old broken behavior too, so a reader recognizes their problem.')

para('At release time', style='Heading 3')
rich('The running file is promoted to ',
     ('docs/release-notes/v<VERSION>.md', {'mono': True}),
     ' and committed. It is then used verbatim for three audiences: the appcast '
     'description shown in the in-app update dialog, the GitHub release body, '
     'and the website changelog entry.')
rich(('Publishing fails if no notes exist.', {'bold': True}),
     ' This is deliberate — a release with no notes cannot be announced, so the '
     'notes cannot fall behind the releases.')
rich('After a release, recreate ', ('next.md', {'mono': True}),
     ' with the same convention header, ready for the next cycle.')

# ======================= 5. THE PIPELINE =======================
para('5. The release pipeline', style='Heading 2')
rich('One command: ', ('NOTARY_PROFILE=sealshot-notary ./scripts/release.sh',
                       {'mono': True}))
para('It runs in stages and stops at the first failure, so a partial release is '
     'never published.')

para('Preflight — before anything is built', style='Heading 3')
numbered('A “Developer ID Application” certificate is present in the keychain.')
numbered('Notarization credentials are available.')
numbered('The GitHub CLI is installed and authenticated.')
numbered('The Sparkle EdDSA signing key is in the keychain.')
numbered('Release notes exist.')
para('Failing early matters: notarization takes minutes, and discovering a '
     'missing credential afterwards wastes the whole run.',
     italic=True, color=MUTED, size=9.5)

para('Build and sign', style='Heading 3')
numbered('Regenerate the Xcode project.')
numbered('Archive a universal binary — Apple silicon and Intel — from the direct-download scheme.')
numbered('Export the signed app and verify the signature.')
numbered('Stamp the entitlement date into the bundle.')
numbered('Build the styled .dmg.')

para('Notarize', style='Heading 3')
numbered('Submit the .dmg to Apple’s notary service and wait for the result.')
numbered('Staple the notarization ticket to the .dmg.')
para('Stapling is what lets the app open with no Gatekeeper warning even offline.')

para('Publish', style='Heading 3')
numbered('Create the update .zip and sign it with the Sparkle key.')
numbered('Promote and commit the release notes.')
numbered('Append an appcast entry, newest first, and validate the XML.')
numbered('Create the GitHub release with the .zip and .dmg attached.')
numbered('Push the updated appcast.')
numbered('Tag the commit.')
numbered('Open a changelog pull request against the website.')

para('The appcast entry carries the build number, the marketing version, the '
     'minimum macOS version, the release notes as HTML, and the signature of '
     'the .zip. An entry for a build already present is skipped rather than '
     'duplicated, and malformed XML aborts the release — an invalid appcast '
     'would break updates for every existing install.')

para('The website changelog pull request is the one non-fatal step: by the time '
     'it runs the release is already public, so a failure there means adding the '
     'changelog entry by hand rather than re-releasing.')

para('Dry runs', style='Heading 3')
table(
    ['Flag', 'Effect', 'Use it for'],
    [
        ['SKIP_NOTARIZE=1', 'Build, sign, and package only', 'Smoke-testing the build without waiting on Apple'],
        ['SKIP_PUBLISH=1', 'Everything except zip, appcast, and GitHub release', 'Handing a signed build to a tester privately'],
    ],
    widths=[1.45, 2.55, 2.65], mono_cols=(0,),
)

# ======================= 6. HOW UPDATES REACH USERS =======================
para('6. How updates reach users', style='Heading 2')
para('Direct downloads update themselves through Sparkle. The app checks a '
     'public appcast once a day, compares build numbers, and offers any newer '
     'release with its notes shown in the dialog. Update checks are the app’s '
     'only routine network activity, and can be turned off in Settings.')
para('Each update .zip is signed with a key whose public half ships inside the '
     'app, so an update that was not built by us is rejected before it is '
     'installed.')

para('Every release stays available', style='Heading 3')
para('Past releases are never removed. Two reasons: a customer whose update '
     'window has lapsed can always reinstall the newest release their license '
     'covers, and anyone who hits a regression can step back a version without '
     'asking for help.')

para('One channel, one build', style='Heading 3')
para('Sealshot ships only as a signed, notarized direct download. There is no '
     'Mac App Store edition, so there is no feature that some customers have and '
     'others do not, and no second entitlement model to reason about.')
table(
    ['', 'How it works'],
    [
        ['Updates', 'Sparkle, from the appcast in the release repo'],
        ['Licensing', 'A signed license file carrying an update window'],
        ['Trial', '14 days, every feature unlocked, no account'],
        ['Entitlement date', 'Stamped into the build at release time'],
        ['Revocation', 'A signed blocklist the app downloads and checks locally'],
    ],
    widths=[1.9, 4.6],
)
para('That the trial, the update window and revocation all resolve on the '
     'customer\u2019s own Mac is the point rather than a limitation: it is what lets '
     'the app promise that nothing leaves the device. The cost is that '
     'revocation is best-effort — see \u00a79.')

# ======================= 7. WEBSITE =======================
para('7. The website side of a release', style='Heading 2')
para('Each release adds one changelog page, created automatically by the '
     'pull request the pipeline opens:')
code_block([
    '---',
    'title: Sealshot 1.0.0',
    'description: Release notes for Sealshot 1.0.0.',
    'slug: docs/changelog/v1-0-0',
    'sidebar:',
    '  order: -1000',
    '---',
])
bullet(('The slug replaces dots with hyphens', {'bold': True}),
       ' — dots in a URL segment would be read as a file extension.')
bullet(('sidebar.order is negative and decreasing', {'bold': True}),
       ' so newer releases sort above older ones automatically.')
rich('Two links elsewhere on the site point at the newest release and need '
     'updating when a release lands: the “What’s new” line on the documentation '
     'landing page, and the changelog link on the marketing home page. The '
     'previous release’s entry should also stop describing itself as the latest.')
para('The download page resolves the current .dmg from the release repository, '
     'so it needs no change per release.')

# ======================= 8. CADENCE =======================
para('8. Cadence and scope', style='Heading 2')
para('Releases are cut when something is worth shipping rather than on a '
     'calendar. In practice that has meant a release every one to three weeks, '
     'mixing features and fixes.')
para('Two rules keep that sustainable:')
bullet(('Never ship without notes.', {'bold': True}),
       ' Enforced by the pipeline, and it is what keeps the changelog honest.')
bullet(('Never ship an unverified build.', {'bold': True}),
       ' Signature verification and notarization are part of the pipeline, not '
       'a manual afterthought.')

# ======================= 9. WITHDRAWING =======================
para('9. Withdrawing a release or a license', style='Heading 2')
para('Two independent mechanisms, for two different problems.')

para('A bad release', style='Heading 3')
para('Remove or replace its appcast entry so it stops being offered, then '
     'publish a fixed release with a higher build number. Existing installs are '
     'unaffected — they only ever move forward. The bad release stays '
     'downloadable unless deliberately removed, which is usually the right '
     'choice: someone already running it may need to reinstall it.')

para('A revoked license — after a refund, or fraud', style='Heading 3')
para('Licenses are revoked through a signed blocklist published alongside the '
     'releases, at license-blocklist.json. The app downloads the whole list and '
     'checks it locally, so revocation needs no new app version and no server is '
     'ever asked about a particular license.')
rich(('    licensegen revoke --id <uuid> --blocklist license-blocklist.json',
      {'mono': True}), space_after=4)
para('Commit the result, then verify what was actually published:')
rich(('    npm run verify:blocklist -- --expect <uuid>', {'mono': True}), space_after=6)
para('That second step is not optional politeness. The app fails open on '
     'purpose — an unreachable or unverifiable blocklist must never stop a '
     'paying customer working — so a 404, a malformed file, or an id added by '
     'hand without re-signing all degrade silently to "revokes nothing". That is '
     'indistinguishable from "nothing has been revoked yet", which is exactly '
     'how this URL returned 404 for months without anyone noticing.')
para('Revocation is not instant, and not uniform. The blocklist is served from '
     'a CDN with many independent edges, each caching for about five minutes, so '
     'for a few minutes after publishing some installs will see the revocation '
     'and others will not. Verify against the GitHub API rather than the raw URL: '
     'the API is uncached, while two raw requests from the same machine can '
     'legitimately disagree.')
para('Two consequences worth being clear-eyed about. Revocation cannot be '
     'guaranteed: a customer who stays offline never receives it, and the cached '
     'list is a file they can delete. And an older, validly signed list still '
     'verifies, so the scheme has no rollback protection. Both are accepted '
     'costs of verifying licenses entirely on the customer\u2019s own machine. '
     'Revocation is an honesty mechanism, not a lock — what genuinely deters '
     'sharing is that the buyer\u2019s name and email are in the file, signed, and '
     'cannot be edited out.')

para('A compromised signing key', style='Heading 3')
para('Two license-signing public keys ship in the app: a primary and a standby. '
     'If the primary is ever exposed, licenses can be issued under the standby '
     'without stranding installs that are already out in the world.')

# ======================= 10. CHECKLIST =======================
para('10. Release checklist', style='Heading 2')
para('Before running the pipeline:')
numbered('Release notes for every user-visible change are in the running notes file.')
numbered('The marketing version is bumped, and the build number is higher than the last published one.')
numbered('Tests pass.')
numbered('A dry run has been built and opened, with notarization skipped.')
para('After it completes:')
numbered('The update dialog offers the new version in an older install, with the notes shown correctly.')
numbered('The downloaded .dmg opens with no Gatekeeper warning on a clean machine.')
numbered('The GitHub release has both the .zip and the .dmg attached.')
numbered('The website changelog pull request is merged, and the “what’s new” links point at it.')
numbered('The running notes file is recreated, empty, for the next cycle.')

out = '/Users/ledeng/projects/Sealshot-Website/output/Sealshot-Release-Cycle-v1.0.docx'
doc.save(out)
print('wrote', out)
